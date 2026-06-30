"""
PDF & DOC -> Markdown Converter
Usage:
    python converter.py <file_or_folder> [options]

Options:
    -o, --output <dir>    Output directory (default: same as input)
    -r, --recursive       Recurse into subdirectories
    --overwrite           Overwrite existing .md files
"""

import argparse
import sys
import os
import re
from pathlib import Path


# ---------------------------------------------------------------------------
# PDF conversion (uses PyMuPDF / fitz)
# ---------------------------------------------------------------------------

def pdf_to_markdown(pdf_path: Path) -> str:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        sys.exit(
            "[ERROR] PyMuPDF is not installed.\n"
            "  Run:  pip install pymupdf"
        )

    doc = fitz.open(str(pdf_path))
    parts: list[str] = []

    for page_num, page in enumerate(doc, start=1):
        # Extract structured blocks: (x0,y0,x1,y1, text, block_no, block_type)
        blocks = page.get_text("dict")["blocks"]

        parts.append(f"\n\n<!-- Page {page_num} -->\n")

        for block in blocks:
            if block.get("type") != 0:   # 0 = text block; skip images etc.
                continue
            for line in block.get("lines", []):
                line_text_parts = []
                max_size = 0.0
                bold_flags = []
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if not text:
                        continue
                    size = span.get("size", 12)
                    flags = span.get("flags", 0)  # bit 4 = bold, bit 1 = italic
                    max_size = max(max_size, size)
                    is_bold = bool(flags & 0b10000)
                    is_italic = bool(flags & 0b00010)
                    if is_bold and is_italic:
                        text = f"***{text}***"
                    elif is_bold:
                        text = f"**{text}**"
                    elif is_italic:
                        text = f"*{text}*"
                    line_text_parts.append(text)
                    bold_flags.append(is_bold)

                line_text = " ".join(line_text_parts).strip()
                if not line_text:
                    continue

                # Heuristic heading detection based on font size
                if max_size >= 20:
                    parts.append(f"\n# {_strip_md_inline(line_text)}\n")
                elif max_size >= 16:
                    parts.append(f"\n## {_strip_md_inline(line_text)}\n")
                elif max_size >= 13:
                    parts.append(f"\n### {_strip_md_inline(line_text)}\n")
                else:
                    parts.append(line_text + "\n")

    doc.close()
    return _clean_markdown("".join(parts))


# ---------------------------------------------------------------------------
# DOCX conversion (uses python-docx)
# ---------------------------------------------------------------------------

def docx_to_markdown(docx_path: Path) -> str:
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        sys.exit(
            "[ERROR] python-docx is not installed.\n"
            "  Run:  pip install python-docx"
        )

    doc = Document(str(docx_path))
    parts: list[str] = []

    # Heading style name → markdown prefix
    HEADING_MAP = {
        "heading 1": "#",
        "heading 2": "##",
        "heading 3": "###",
        "heading 4": "####",
        "heading 5": "#####",
        "heading 6": "######",
    }

    # Track list state
    list_counter: dict[int, int] = {}  # ilvl → count (for numbered lists)
    prev_was_list = False

    for para in doc.paragraphs:
        style_name = para.style.name.lower()
        text = _runs_to_md(para.runs)

        if not text.strip():
            if prev_was_list:
                parts.append("\n")
                prev_was_list = False
            else:
                parts.append("\n")
            continue

        # Headings
        if style_name in HEADING_MAP:
            prefix = HEADING_MAP[style_name]
            parts.append(f"\n{prefix} {text.strip()}\n\n")
            prev_was_list = False
            list_counter.clear()
            continue

        # Bullet / numbered lists via paragraph format
        num_pr = para._p.find(qn("w:numPr"))
        if num_pr is not None:
            ilvl_el = num_pr.find(qn("w:ilvl"))
            num_id_el = num_pr.find(qn("w:numId"))
            ilvl = int(ilvl_el.get(qn("w:val"), 0)) if ilvl_el is not None else 0
            num_id = int(num_id_el.get(qn("w:val"), 0)) if num_id_el is not None else 0
            indent = "  " * ilvl

            # Try to detect ordered vs unordered from numbering definition
            is_ordered = _is_ordered_list(doc, num_id, ilvl)
            if is_ordered:
                list_counter[ilvl] = list_counter.get(ilvl, 0) + 1
                # reset deeper levels
                for deeper in list(list_counter.keys()):
                    if deeper > ilvl:
                        del list_counter[deeper]
                parts.append(f"{indent}{list_counter[ilvl]}. {text.strip()}\n")
            else:
                parts.append(f"{indent}- {text.strip()}\n")
            prev_was_list = True
            continue

        prev_was_list = False

        # Block-quote style
        if "quote" in style_name:
            parts.append(f"\n> {text.strip()}\n\n")
            continue

        # Code / preformatted
        if "code" in style_name or "mono" in style_name:
            parts.append(f"\n```\n{text.strip()}\n```\n\n")
            continue

        # Normal paragraph
        parts.append(f"\n{text.strip()}\n")

    # Tables
    for table in doc.tables:
        parts.append(_table_to_md(table))

    return _clean_markdown("".join(parts))


# ---------------------------------------------------------------------------
# DOC (legacy binary) – requires LibreOffice or antiword as fallback
# ---------------------------------------------------------------------------

def doc_to_markdown(doc_path: Path) -> str:
    """Convert legacy .doc file by first converting to .docx via LibreOffice."""
    import subprocess
    import tempfile

    # Try LibreOffice headless conversion
    lo_candidates = [
        "libreoffice",
        "soffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]

    lo_bin = None
    for candidate in lo_candidates:
        try:
            result = subprocess.run(
                [candidate, "--version"],
                capture_output=True, text=True, timeout=10
            )
            if result.returncode == 0:
                lo_bin = candidate
                break
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    if lo_bin is None:
        sys.exit(
            "[ERROR] Legacy .doc files require LibreOffice.\n"
            "  Install from https://www.libreoffice.org/ and ensure 'soffice' is on PATH.\n"
            "  Alternatively, open the file in Word and save as .docx first."
        )

    with tempfile.TemporaryDirectory() as tmpdir:
        subprocess.run(
            [lo_bin, "--headless", "--convert-to", "docx",
             "--outdir", tmpdir, str(doc_path)],
            check=True, capture_output=True
        )
        docx_files = list(Path(tmpdir).glob("*.docx"))
        if not docx_files:
            sys.exit(f"[ERROR] LibreOffice failed to convert {doc_path}")
        return docx_to_markdown(docx_files[0])


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _runs_to_md(runs) -> str:
    """Convert docx Run objects to inline Markdown."""
    out = []
    for run in runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        elif run.underline:
            # Markdown has no underline; use HTML or just leave plain
            text = f"<u>{text}</u>"
        if run.font.strike:
            text = f"~~{text}~~"
        out.append(text)
    return "".join(out)


def _is_ordered_list(doc, num_id: int, ilvl: int) -> bool:
    """Check numbering definition to decide ordered vs unordered."""
    try:
        from docx.oxml.ns import qn
        numbering = doc.part.numbering_part
        if numbering is None:
            return False
        root = numbering._element
        # Find abstractNumId for this numId
        for num_el in root.findall(qn("w:num")):
            if num_el.get(qn("w:numId")) == str(num_id):
                abstract_ref = num_el.find(qn("w:abstractNumId"))
                if abstract_ref is None:
                    return False
                abstract_id = abstract_ref.get(qn("w:val"))
                # Find the abstractNum
                for abs_num in root.findall(qn("w:abstractNum")):
                    if abs_num.get(qn("w:abstractNumId")) == abstract_id:
                        for lvl in abs_num.findall(qn("w:lvl")):
                            if lvl.get(qn("w:ilvl")) == str(ilvl):
                                num_fmt = lvl.find(qn("w:numFmt"))
                                if num_fmt is not None:
                                    fmt = num_fmt.get(qn("w:val"), "bullet")
                                    return fmt not in ("bullet", "none")
        return False
    except Exception:
        return False


def _table_to_md(table) -> str:
    """Convert a docx Table to a Markdown table."""
    rows = table.rows
    if not rows:
        return ""

    md_rows = []
    for i, row in enumerate(rows):
        cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
        md_rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            md_rows.append("| " + " | ".join(["---"] * len(cells)) + " |")

    return "\n\n" + "\n".join(md_rows) + "\n\n"


def _strip_md_inline(text: str) -> str:
    """Remove inline markdown markers from a heading candidate."""
    return re.sub(r"[*_~`]", "", text)


def _clean_markdown(text: str) -> str:
    """Normalize whitespace and fix common conversion artifacts."""
    # Collapse 3+ blank lines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove trailing spaces
    text = re.sub(r"[ \t]+\n", "\n", text)
    return text.strip() + "\n"


# ---------------------------------------------------------------------------
# File dispatch
# ---------------------------------------------------------------------------

def convert_file(input_path: Path, output_dir: Path | None, overwrite: bool) -> None:
    suffix = input_path.suffix.lower()

    if suffix == ".pdf":
        md_content = pdf_to_markdown(input_path)
    elif suffix == ".docx":
        md_content = docx_to_markdown(input_path)
    elif suffix == ".doc":
        md_content = doc_to_markdown(input_path)
    else:
        print(f"[SKIP] Unsupported format: {input_path}")
        return

    out_dir = output_dir if output_dir else input_path.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / (input_path.stem + ".md")

    if out_path.exists() and not overwrite:
        print(f"[SKIP] Already exists (use --overwrite): {out_path}")
        return

    out_path.write_text(md_content, encoding="utf-8")
    print(f"[OK]   {input_path.name}  ->  {out_path}")


def collect_files(path: Path, recursive: bool) -> list[Path]:
    SUPPORTED = {".pdf", ".doc", ".docx"}
    if path.is_file():
        return [path] if path.suffix.lower() in SUPPORTED else []
    pattern = "**/*" if recursive else "*"
    return [p for p in path.glob(pattern) if p.suffix.lower() in SUPPORTED]


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convert PDF / DOC / DOCX files to Markdown (.md)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument(
        "input",
        help="Path to a file or folder containing PDF/DOC/DOCX files.",
    )
    parser.add_argument(
        "-o", "--output",
        metavar="DIR",
        help="Directory to write .md files into (default: same as input).",
    )
    parser.add_argument(
        "-r", "--recursive",
        action="store_true",
        help="Recursively search subdirectories.",
    )
    parser.add_argument(
        "--overwrite",
        action="store_true",
        help="Overwrite existing .md files.",
    )

    args = parser.parse_args()
    input_path = Path(args.input).expanduser().resolve()

    if not input_path.exists():
        sys.exit(f"[ERROR] Path not found: {input_path}")

    output_dir = Path(args.output).expanduser().resolve() if args.output else None

    files = collect_files(input_path, args.recursive)
    if not files:
        print("[WARN] No PDF / DOC / DOCX files found.")
        return

    print(f"Found {len(files)} file(s) to convert.\n")
    for f in files:
        try:
            convert_file(f, output_dir, args.overwrite)
        except Exception as exc:
            print(f"[ERROR] {f.name}: {exc}")

    print("\nDone.")


if __name__ == "__main__":
    main()
